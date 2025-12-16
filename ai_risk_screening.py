
import streamlit as st
import pandas as pd
import json
import re
from io import BytesIO

# -------------------- Page config & simple styling --------------------
st.set_page_config(page_title="AI Risk & Criticality Screening", layout="wide")
st.markdown(
    """
    <style>
    .section-title {font-size:1.25rem; font-weight:700; margin-top:1.0rem; padding-bottom:0.3rem; border-bottom:1px solid #ddd;}
    .small-note {font-size:0.9rem; color:#666;}
    </style>
    """,
    unsafe_allow_html=True,
)
st.title("AI Use Case Risk & Criticality Screening Questionnaire")

# -------------------- Helpers: core logic --------------------
def enrich_with_other(selected_list, other_text):
    """If the user chose 'Other', replace it with comma-separated values from the text box."""
    if "Other" in selected_list:
        selected_list = [s for s in selected_list if s != "Other"]
        if other_text and other_text.strip():
            extras = [x.strip() for x in other_text.split(",") if x.strip()]
            selected_list.extend(extras)
    return selected_list

def max_source_risk(sources):
    """Map source types to risk levels and return the maximum."""
    mapping = {"Internal": 2, "External": 3, "Third-party": 4}
    if not sources:
        return 3
    return max([mapping.get(s, 3) for s in sources if s in mapping])

def regulatory_base_risk(regs):
    """Increase risk if certain regulations apply; otherwise use presence/absence as a proxy."""
    regulated = {"GDPR", "HIPAA", "RBI"}
    if any(r in regulated for r in regs):
        return 4
    if len(regs) > 0:
        return 3
    return 2

def categorize_total_risk(total):
    """Categorize weighted total into Low/Medium/High."""
    if total >= 3.5:
        return "High"
    elif total >= 2.5:
        return "Medium"
    else:
        return "Low"

def decision_matrix(risk_level, criticality):
    """Simple decision matrix for recommendations."""
    if risk_level == "Low":
        return "Approve"
    if risk_level == "Medium" and criticality == "High":
        return "Approve with conditions"
    if risk_level == "Medium":
        return "Approve"
    if risk_level == "High" and criticality == "High":
        return "Escalate for detailed review"
    if risk_level == "High":
        return "Reject or redesign"
    return "Review"

def performance_metrics_risk(metrics):
    """Heuristic: only Accuracy is weak; presence of standard metrics is stronger."""
    if not metrics:
        return 4.5
    std = {"Accuracy", "Precision", "Recall", "F1 Score"}
    only_accuracy = (metrics == ["Accuracy"])
    has_std = any(m in std for m in metrics)
    if only_accuracy:
        return 3.5
    if has_std:
        return 2.0
    return 3.0

# -------------------- Abbreviation expansion (labels + values + free-text) --------------------
ABBREV_MAP_VALUES = {
    "NLP": "Natural Language Processing",
    "CSAT": "Customer Satisfaction",
    "RBAC": "Role-Based Access Control",
    "API": "Application Programming Interface",
    "On-prem": "On-premises",
}

ABBREV_MAP_LABELS = {
    "Use Case Desc": "Use Case Description",
    "Deployment Envs": "Deployment Environments",
    "Regs": "Applicable Regulations",
    "Bias Fairness": "Bias and Fairness Checks",
    "Privacy By Design": "Privacy by Design",
    "Cyber Measures": "Cybersecurity Measures",
}

# Regex patterns for expansions (case-insensitive)
TEXT_ABBREV_PATTERNS = [
    (r"\bCSAT\b", "Customer Satisfaction"),
    (r"\bNLP\b", "Natural Language Processing"),
    (r"\bRBAC\b", "Role-Based Access Control"),
    (r"\bAPI\b", "Application Programming Interface"),
    (r"\bDLP\b", "Data Loss Prevention"),
    (r"\bDPIA\b", "Data Protection Impact Assessment"),
    (r"\bMFA\b", "Multi-Factor Authentication"),
    (r"\bIR\b", "Incident Response"),
    (r"\bPII\b", "Personally Identifiable Information"),
    (r"\bPHI\b", "Protected Health Information"),
    (r"\bGDPR\b", "General Data Protection Regulation"),
    (r"\bHIPAA\b", "Health Insurance Portability and Accountability Act"),
    (r"\bRBI\b", "Reserve Bank of India regulations"),
    (r"\bEU AI Act\b", "European Union Artificial Intelligence Act"),
    (r"\bISO(?:\/IEC)?\s*42001\b", "ISO/IEC 42001 Artificial Intelligence Management System standard"),
    (r"\bNIST\s*AI\s*RMF\b", "NIST AI Risk Management Framework"),
]

def expand_text(s: str) -> str:
    """
    Expand acronyms/short forms in free text:
      - Case-insensitive regex replacements with word boundaries.
      - Common shorthand normalization (e.g., 24x7 -> 24 hours a day, 7 days a week).
    """
    if not isinstance(s, str):
        return s
    expanded = s
    for pat, repl in TEXT_ABBREV_PATTERNS:
        expanded = re.sub(pat, repl, expanded, flags=re.IGNORECASE)
    # Shorthand/symbols:
    expanded = re.sub(r"\b24\s*[x×]\s*7\b", "24 hours a day, 7 days a week", expanded, flags=re.IGNORECASE)
    expanded = re.sub(r"\b24\s*/\s*7\b", "24 hours a day, 7 days a week", expanded, flags=re.IGNORECASE)
    return expanded

def expand_label(label: str) -> str:
    """Expand shorthand in labels/headings."""
    return ABBREV_MAP_LABELS.get(label, label)

def expand_value(val):
    """Expand acronyms in values (string or list) and run free-text expansion."""
    if isinstance(val, str):
        v = ABBREV_MAP_VALUES.get(val, val)
        return expand_text(v)
    if isinstance(val, list):
        return [expand_text(ABBREV_MAP_VALUES.get(v, v)) for v in val]
    return val

def titleize_key(k: str) -> str:
    """Convert snake_case keys to Title Case and expand label shorthand."""
    pretty = k.replace("_", " ").strip().title()
    return expand_label(pretty)

def expand_dict_keys_values(d: dict) -> dict:
    """Produce a dict with expanded labels and values for output files."""
    out = {}
    for k, v in d.items():
        key_expanded = titleize_key(k)
        out[key_expanded] = expand_value(v)
    return out

# -------------------- Section 1: Use Case Overview --------------------
st.markdown('<div class="section-title">Section 1: Use Case Overview</div>', unsafe_allow_html=True)
c1, c2 = st.columns(2)
with c1:
    use_case_name = st.text_input("Use Case Name")
with c2:
    use_case_desc = st.text_area("Use Case Description")

business_objective = st.text_area("Business Objective")

model_types_base = ["Predictive", "Generative", "Natural Language Processing", "Classification", "Other"]  # expanded instead of "NLP"
model_types_sel = st.multiselect("AI Model Types (select one or more)", model_types_base)
model_types_other = st.text_input("If 'Other', specify AI Model Types (comma-separated)") if "Other" in model_types_sel else ""
model_types = enrich_with_other(model_types_sel, model_types_other)

deployment_envs_base = ["On-prem", "Cloud", "Hybrid", "Other"]
deployment_envs_sel = st.multiselect("Deployment Environments (select one or more)", deployment_envs_base)
deployment_envs_other = st.text_input("If 'Other', specify Deployment Environments (comma-separated)") if "Other" in deployment_envs_sel else ""
deployment_envs = enrich_with_other(deployment_envs_sel, deployment_envs_other)

# -------------------- Section 2: Data Risk Assessment --------------------
st.markdown('<div class="section-title">Section 2: Data Risk Assessment</div>', unsafe_allow_html=True)
data_sources_base = ["Internal", "External", "Third-party", "Other"]
data_sources_sel = st.multiselect("Data Sources (select one or more)", data_sources_base)
data_sources_other = st.text_input("If 'Other', specify Data Sources (comma-separated)") if "Other" in data_sources_sel else ""
data_sources = enrich_with_other(data_sources_sel, data_sources_other)

data_sensitivity = st.radio(
    "Data Sensitivity (Personally Identifiable Information / Protected Health Information / Financial / Confidential) present?",
    ["Yes", "No"],
    horizontal=True,
)
data_bias_checks = st.radio("Bias and Quality checks implemented?", ["Yes", "No"], horizontal=True)
data_encryption = st.radio("Encryption and retention policies applied?", ["Yes", "No"], horizontal=True)

# -------------------- Section 3: Model Risk Assessment --------------------
st.markdown('<div class="section-title">Section 3: Model Risk Assessment</div>', unsafe_allow_html=True)
model_explainable = st.radio("Model transparency (explainable)?", ["Yes", "No"], horizontal=True)
bias_fairness = st.radio("Bias and Fairness Checks implemented?", ["Yes", "No"], horizontal=True)

perf_options = ["Accuracy", "Precision", "Recall", "F1 Score", "Other"]
performance_metrics_sel = st.multiselect("Performance Metrics (select one or more)", perf_options)
performance_metrics_other = st.text_input("If 'Other', specify Performance Metrics (comma-separated)") if "Other" in performance_metrics_sel else ""
performance_metrics = enrich_with_other(performance_metrics_sel, performance_metrics_other)

model_drift = st.radio("Model drift monitoring in place?", ["Yes", "No"], horizontal=True)

# -------------------- Section 4: Operational Risk --------------------
st.markdown('<div class="section-title">Section 4: Operational Risk</div>', unsafe_allow_html=True)
criticality = st.selectbox("Criticality of impacted business process", ["Low", "Medium", "High"])
dependency = st.selectbox("Dependency on AI output", ["Advisory", "Fully Automated"])
fallback = st.radio("Fallback mechanism / human-in-the-loop available?", ["Yes", "No"], horizontal=True)

# -------------------- Section 5: Compliance & Regulatory --------------------
st.markdown('<div class="section-title">Section 5: Compliance & Regulatory</div>', unsafe_allow_html=True)
regs_base = ["GDPR", "HIPAA", "RBI", "Other"]
regs_sel = st.multiselect("Applicable Regulations (select one or more)", regs_base)
regs_other = st.text_input("If 'Other', specify applicable regulations (comma-separated)") if "Other" in regs_sel else ""
regs = enrich_with_other(regs_sel, regs_other)

consent = st.radio("Consent management present?", ["Yes", "No"], horizontal=True)
auditability = st.radio("Decisions traceable/auditable?", ["Yes", "No"], horizontal=True)

# -------------------- Section 6: Security & Privacy --------------------
st.markdown('<div class="section-title">Section 6: Security & Privacy</div>', unsafe_allow_html=True)
access_controls = st.radio("Role-based access controls?", ["Yes", "No"], horizontal=True)
cyber_measures = st.radio("Cybersecurity measures implemented?", ["Yes", "No"], horizontal=True)
privacy_by_design = st.radio("Privacy by Design embedded?", ["Yes", "No"], horizontal=True)

# -------------------- Section 7: Risk & Criticality Scoring (internal only) --------------------
st.markdown('<div class="section-title">Section 7: Risk & Criticality Scoring</div>', unsafe_allow_html=True)
st.markdown('<p class="small-note">Weights are 20% for Data, Model, Operational, Regulatory, Security. 1=low risk, 5=high risk.</p>', unsafe_allow_html=True)

data_risk_components = [
    5 if data_sensitivity == "Yes" else 2,
    4 if data_bias_checks == "No" else 2,
    4 if data_encryption == "No" else 2,
    max_source_risk([s for s in data_sources if s in ["Internal", "External", "Third-party"]]),
]
data_risk_score = round(sum(data_risk_components) / len(data_risk_components), 2)

model_risk_components = [
    4 if model_explainable == "No" else 2,
    4 if bias_fairness == "No" else 2,
    performance_metrics_risk(performance_metrics),
    4 if model_drift == "No" else 2,
]
model_risk_score = round(sum(model_risk_components) / len(model_risk_components), 2)

crit_map = {"Low": 2, "Medium": 3, "High": 4}
dep_map = {"Advisory": 2, "Fully Automated": 4}
operational_risk_components = [
    crit_map.get(criticality, 3),
    dep_map.get(dependency, 3),
    4 if fallback == "No" else 2,
]
operational_risk_score = round(sum(operational_risk_components) / len(operational_risk_components), 2)

reg_base = regulatory_base_risk(regs)
regulatory_risk_components = [
    reg_base,
    4 if consent == "No" else 2,
    4 if auditability == "No" else 2,
]
regulatory_risk_score = round(sum(regulatory_risk_components) / len(regulatory_risk_components), 2)

security_risk_components = [
    4 if access_controls == "No" else 2,
    4 if cyber_measures == "No" else 2,
    4 if privacy_by_design == "No" else 2,
]
security_risk_score = round(sum(security_risk_components) / len(security_risk_components), 2)

weights = {"Data": 0.20, "Model": 0.20, "Operational": 0.20, "Regulatory": 0.20, "Security": 0.20}
total_weighted = round(
    data_risk_score * weights["Data"]
    + model_risk_score * weights["Model"]
    + operational_risk_score * weights["Operational"]
    + regulatory_risk_score * weights["Regulatory"]
    + security_risk_score * weights["Security"],
    2,
)
risk_level = categorize_total_risk(total_weighted)
recommendation = decision_matrix(risk_level, criticality)

# -------------------- Submit & Results (expanded outputs + downloads) --------------------
if st.button("Submit & Analyze"):
    payload = {
        "use_case_name": use_case_name,
        "use_case_desc": use_case_desc,
        "business_objective": business_objective,
        "model_types": model_types,
        "deployment_envs": deployment_envs,
        "data_sources": data_sources,
        "data_sensitivity": data_sensitivity,
        "data_bias_checks": data_bias_checks,
        "data_encryption": data_encryption,
        "model_explainable": model_explainable,
        "bias_fairness": bias_fairness,
        "performance_metrics": performance_metrics,
        "model_drift": model_drift,
        "criticality": criticality,
        "dependency": dependency,
        "fallback": fallback,
        "regs": regs,
        "consent": consent,
        "auditability": auditability,
        "access_controls": access_controls,
        "cyber_measures": cyber_measures,
        "privacy_by_design": privacy_by_design,
        "scores": {
            "data_risk": data_risk_score,
            "model_risk": model_risk_score,
            "operational_risk": operational_risk_score,
            "regulatory_risk": regulatory_risk_score,
            "security_risk": security_risk_score,
            "total_weighted": total_weighted,
            "risk_level": risk_level,
            "recommendation": recommendation,
        },
    }

    # Prepare user input for output (no scores, no identified risks)
    user_input_raw = {k: v for k, v in payload.items() if k not in {"scores", "identified_risks"}}

    # Expand free-text and values for display/download
    user_input_for_output = {}
    for k, v in user_input_raw.items():
        if isinstance(v, str):
            user_input_for_output[k] = expand_text(v)
        elif isinstance(v, list):
            user_input_for_output[k] = [expand_text(x) for x in v]
        else:
            user_input_for_output[k] = v

    expanded_user_input = expand_dict_keys_values(user_input_for_output)

    st.success("Submission captured. Download your responses below (DOCX / PDF / XLSX / JSON).")
    st.subheader("User Input (Expanded)")
    st.json(expanded_user_input)

    # -------------------- File builders (DOCX, PDF, XLSX, JSON) --------------------
    def build_docx(data: dict) -> bytes:
        from docx import Document

        doc = Document()
        doc.add_heading("AI Use Case Risk & Criticality – User Submission", level=1)

        for k, v in data.items():
            p = doc.add_paragraph()
            p.add_run(f"{k}: ").bold = True
            p.add_run(", ".join(v) if isinstance(v, list) else str(v))

        bio = BytesIO()
        doc.save(bio)
        return bio.getvalue()

    def build_pdf(data: dict) -> bytes:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import cm

        story = []
        styles = getSampleStyleSheet()
        story.append(Paragraph("AI Use Case Risk & Criticality – User Submission", styles["Title"]))
        story.append(Spacer(1, 0.3 * cm))

        for k, v in data.items():
            val = ", ".join(v) if isinstance(v, list) else str(v)
            story.append(Paragraph(f"<b>{k}:</b> {val}", styles["BodyText"]))
            story.append(Spacer(1, 0.2 * cm))

        bio = BytesIO()
        pdf = SimpleDocTemplate(
            bio,
            pagesize=A4,
            leftMargin=2 * cm,
            rightMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
        )
        pdf.build(story)
        return bio.getvalue()

    def build_xlsx(data: dict) -> bytes:
        rows = []
        for k, v in data.items():
            val = ", ".join(v) if isinstance(v, list) else v
            rows.append({"Field": k, "Value": val})

        df = pd.DataFrame(rows)
        bio = BytesIO()
        with pd.ExcelWriter(bio, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name="User Input")
        return bio.getvalue()

    # Build files from expanded output dict
    docx_bytes = build_docx(expanded_user_input)
    pdf_bytes = build_pdf(expanded_user_input)
    xlsx_bytes = build_xlsx(expanded_user_input)
    json_bytes = json.dumps(expanded_user_input, indent=2).encode("utf-8")

    # Download buttons
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.download_button(
            "Download DOCX",
            data=docx_bytes,
            file_name="ai_risk_user_input.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    with c2:
        st.download_button(
            "Download PDF",
            data=pdf_bytes,
            file_name="ai_risk_user_input.pdf",
            mime="application/pdf",
        )
    with c3:
        st.download_button(
            "Download            "Download XLSX",
            data=xlsx_bytes,
            file_name="ai_risk_user_input.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
    with c4:
        st.download_button(
            "Download JSON",
            data=json_bytes,
            file_name="user_input.json",
            mime="application/json",

